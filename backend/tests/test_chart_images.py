from pathlib import Path

from docx import Document
from sqlalchemy import select

from app.core.database import SessionLocal
from app.models.export_artifact import ExportArtifact
from tests.test_analysis_orchestrator import prepare_orchestrator_fixture


PNG_BASE64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="


def _assert_no_exe_or_u_exe() -> None:
    root = Path("E:/Colmena/backend")
    assert not list(root.rglob("u.exe"))
    assert not [path for path in root.rglob("*.exe") if ".venv" not in str(path).lower()]


def test_chart_images_and_word_integration(client):
    fixture = prepare_orchestrator_fixture(client)
    form_id = fixture["form"]["id"]

    upload = client.post(
        f"/api/v1/forms/{form_id}/chart-images",
        json={
            "chart_id": "manual-test",
            "chart_type": "bar",
            "title": "Grafico de prueba",
            "format": "png",
            "data_url": f"data:image/png;base64,{PNG_BASE64}",
            "file_name": "grafico_prueba.png",
            "source_type": "chart_editor",
            "metadata_json": {"theme": "colmena_premium"},
        },
    )
    assert upload.status_code == 201
    upload_body = upload.json()
    artifact_id = upload_body["image"]["artifact_id"]
    file_path = Path(upload_body["image"]["file_path"])
    assert artifact_id
    assert file_path.exists()

    with SessionLocal() as session:
        artifact = session.scalar(select(ExportArtifact).where(ExportArtifact.id == artifact_id))
        assert artifact is not None
        assert artifact.artifact_type == "chart_image_png"

    listing = client.get(f"/api/v1/forms/{form_id}/chart-images")
    assert listing.status_code == 200
    assert listing.json()["total"] >= 1

    word_ready = client.get(f"/api/v1/forms/{form_id}/chart-images/word-ready")
    assert word_ready.status_code == 200
    assert word_ready.json()["total"] >= 1

    image_metadata = client.get(f"/api/v1/forms/{form_id}/chart-images/{artifact_id}")
    assert image_metadata.status_code == 200
    assert image_metadata.json()["format"] == "png"

    image_base64 = client.get(f"/api/v1/forms/{form_id}/chart-images/{artifact_id}/base64")
    assert image_base64.status_code == 200
    base64_body = image_base64.json()
    assert base64_body["artifact_id"] == artifact_id
    assert base64_body["data_base64"] == f"data:image/png;base64,{PNG_BASE64}"

    report_with_images = client.post(
        f"/api/v1/forms/{form_id}/word-reports/generate",
        json={
            "report_type": "full_form_report",
            "source_type": "mixed",
            "include_charts_placeholders": True,
            "include_chart_images": True,
            "chart_image_mode": "images_if_available",
            "include_cover": True,
        },
    )
    assert report_with_images.status_code == 200
    report_with_images_body = report_with_images.json()
    assert Path(report_with_images_body["file_path"]).exists()
    assert report_with_images_body["chart_image_count"] >= 1

    document = Document(report_with_images_body["file_path"])
    visible_text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    assert len(document.paragraphs) > 0
    assert len(document.inline_shapes) >= 1
    assert "None" not in visible_text
    assert "nan" not in visible_text.lower()

    report_with_placeholders = client.post(
        f"/api/v1/forms/{form_id}/word-reports/generate",
        json={
            "report_type": "full_form_report",
            "source_type": "mixed",
            "include_charts_placeholders": True,
            "include_chart_images": True,
            "chart_image_mode": "placeholders_only",
            "include_cover": True,
        },
    )
    assert report_with_placeholders.status_code == 200
    assert report_with_placeholders.json()["chart_placeholder_count"] >= 0

    selected_only = client.post(
        f"/api/v1/forms/{form_id}/word-reports/generate",
        json={
            "report_type": "full_form_report",
            "source_type": "mixed",
            "include_charts_placeholders": True,
            "include_chart_images": True,
            "chart_image_mode": "selected_images_only",
            "chart_image_artifact_ids": [artifact_id],
            "include_cover": True,
        },
    )
    assert selected_only.status_code == 200
    assert selected_only.json()["chart_image_artifact_ids"] == [artifact_id]

    invalid_format = client.post(
        f"/api/v1/forms/{form_id}/chart-images",
        json={
            "format": "jpg",
            "data_url": f"data:image/png;base64,{PNG_BASE64}",
        },
    )
    assert invalid_format.status_code == 422

    invalid_data_url = client.post(
        f"/api/v1/forms/{form_id}/chart-images",
        json={
            "format": "png",
            "data_url": "data:image/png;base64,not-valid-base64",
        },
    )
    assert invalid_data_url.status_code == 400

    other_fixture = prepare_orchestrator_fixture(client)
    other_form_id = other_fixture["form"]["id"]
    other_upload = client.post(
        f"/api/v1/forms/{other_form_id}/chart-images",
        json={
            "chart_id": "foreign-test",
            "chart_type": "bar",
            "title": "Grafico ajeno",
            "format": "png",
            "data_url": f"data:image/png;base64,{PNG_BASE64}",
            "file_name": "grafico_ajeno.png",
            "source_type": "chart_editor",
        },
    )
    assert other_upload.status_code == 201
    foreign_artifact_id = other_upload.json()["image"]["artifact_id"]

    invalid_selected = client.post(
        f"/api/v1/forms/{form_id}/word-reports/generate",
        json={
            "report_type": "full_form_report",
            "source_type": "mixed",
            "include_charts_placeholders": True,
            "include_chart_images": True,
            "chart_image_mode": "selected_images_only",
            "chart_image_artifact_ids": [foreign_artifact_id],
            "include_cover": True,
        },
    )
    assert invalid_selected.status_code == 404

    delete_response = client.delete(f"/api/v1/forms/{form_id}/chart-images/{artifact_id}")
    assert delete_response.status_code == 200
    assert delete_response.json()["status"] == "deleted"
    assert not file_path.exists()

    listing_after_delete = client.get(f"/api/v1/forms/{form_id}/chart-images")
    assert listing_after_delete.status_code == 200
    assert all(item["artifact_id"] != artifact_id for item in listing_after_delete.json()["items"])

    _assert_no_exe_or_u_exe()
