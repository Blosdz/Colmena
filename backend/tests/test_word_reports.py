from pathlib import Path

from docx import Document
from sqlalchemy import func, select

from app.core.database import SessionLocal
from app.models.export_artifact import ExportArtifact
from tests.test_analysis_orchestrator import prepare_orchestrator_fixture


def _count_word_reports() -> int:
    with SessionLocal() as session:
        return int(
            session.scalar(
                select(func.count()).select_from(ExportArtifact).where(ExportArtifact.artifact_type == "word_report_docx")
            )
            or 0
        )


def _visible_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text).strip()


def _assert_no_exe_or_u_exe() -> None:
    root = Path("E:/Colmena/backend")
    assert not list(root.rglob("u.exe"))
    assert not [path for path in root.rglob("*.exe") if ".venv" not in str(path).lower()]


def test_word_reports_end_to_end(client):
    fixture = prepare_orchestrator_fixture(client)
    form_id = fixture["form"]["id"]
    questions = fixture["questions"]

    orchestrated = client.post(
        f"/api/v1/forms/{form_id}/analysis/run",
        json={
            "analysis_goal": "correlation",
            "targets": [
                {"target_type": "question", "target_id": questions["edad"]["id"], "role": "x"},
                {"target_type": "question", "target_id": questions["rendimiento"]["id"], "role": "y"},
            ],
            "method": "auto",
            "store_result": True,
            "options": {"include_normality": True, "include_descriptives": True, "include_recommendations": True},
        },
    )
    assert orchestrated.status_code == 200
    orchestrated_run_id = orchestrated.json()["analysis_run_id"]

    correlation = client.post(
        f"/api/v1/forms/{form_id}/correlations/pair",
        json={
            "x": {"target_type": "question", "target_id": questions["edad"]["id"]},
            "y": {"target_type": "question", "target_id": questions["rendimiento"]["id"]},
            "method": "pearson",
            "store_result": True,
        },
    )
    assert correlation.status_code == 200
    correlation_run_id = correlation.json()["analysis_run_id"]

    reports_before = _count_word_reports()
    full_report = client.post(
        f"/api/v1/forms/{form_id}/word-reports/generate",
        json={
            "report_type": "full_form_report",
            "source_type": "mixed",
            "analysis_run_ids": [],
            "title": "Informe de resultados estadisticos",
            "subtitle": "Reporte generado por Colmena",
            "include_charts_placeholders": True,
            "include_plain_language_explanations": True,
            "include_technical_appendix": False,
            "include_cover": True,
            "include_methodology_summary": True,
            "options": {},
        },
    )
    assert full_report.status_code == 200
    full_body = full_report.json()
    assert full_body["artifact_id"] is not None
    assert full_body["file_size_bytes"] > 0
    assert full_body["table_count"] >= 1
    assert full_body["chart_placeholder_count"] >= 1
    assert _count_word_reports() == reports_before + 1

    full_path = Path(full_body["file_path"])
    assert full_path.exists()
    document = Document(full_path)
    text = _visible_text(document)
    assert "Informe de resultados estadisticos" in text
    assert "Datos generales del estudio" in text
    assert "Resultados descriptivos" in text
    assert "Graficos sugeridos" in text
    assert "Nota." in text
    assert "None" not in text
    assert "nan" not in text.lower()
    assert len(document.tables) >= 1

    from_run = client.post(f"/api/v1/forms/{form_id}/word-reports/from-analysis-run/{correlation_run_id}")
    assert from_run.status_code == 200
    assert from_run.json()["report_type"] == "inferential_report"

    from_orchestrated = client.post(f"/api/v1/forms/{form_id}/word-reports/from-orchestrated-run/{orchestrated_run_id}")
    assert from_orchestrated.status_code == 200
    assert from_orchestrated.json()["report_type"] == "orchestrated_report"

    options = client.get(f"/api/v1/forms/{form_id}/word-reports/options")
    assert options.status_code == 200
    assert "full_form_report" in options.json()["available_report_types"]

    listing = client.get(f"/api/v1/forms/{form_id}/word-reports")
    assert listing.status_code == 200
    assert len(listing.json()) >= 3

    with SessionLocal() as session:
        artifact = session.scalar(
            select(ExportArtifact)
            .where(ExportArtifact.artifact_type == "word_report_docx")
            .order_by(ExportArtifact.created_at.desc())
        )
        assert artifact is not None

    _assert_no_exe_or_u_exe()


def test_word_report_without_responses_warns(client):
    project_response = client.post("/api/v1/projects", json={"title": "Proyecto vacio"})
    assert project_response.status_code == 201
    project = project_response.json()
    form_response = client.post(f"/api/v1/projects/{project['id']}/forms", json={"title": "Formulario vacio"})
    assert form_response.status_code == 201
    form = form_response.json()

    report = client.post(
        f"/api/v1/forms/{form['id']}/word-reports/generate",
        json={
            "report_type": "full_form_report",
            "source_type": "mixed",
            "include_cover": True,
            "include_charts_placeholders": True,
        },
    )
    assert report.status_code == 200
    body = report.json()
    assert "no_responses" in body["warnings"]
    assert Path(body["file_path"]).exists()
    _assert_no_exe_or_u_exe()
