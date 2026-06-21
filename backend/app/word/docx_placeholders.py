from __future__ import annotations

from pathlib import Path

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _add_cell_shading(cell, fill_hex: str = "E8E8E8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def _set_row_height(row, height_twips: int) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_height)


def _set_table_width(table, width_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def _placeholder_box(document, *, label: str, chart_type: str, explanation: str) -> None:
    """Inserts a gray bordered box as a visual chart placeholder."""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_borders(table)
    _set_table_width(table, 8640)  # 6 inches in twips (1440 twips/inch)

    row = table.rows[0]
    _set_row_height(row, 3960)  # 2.75 inches

    cell = row.cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _add_cell_shading(cell, "F0F0F0")

    label_para = cell.paragraphs[0]
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_para.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if chart_type:
        type_para = cell.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_run = type_para.add_run(f"Tipo sugerido: {chart_type}")
        type_run.font.size = Pt(9)
        type_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if explanation:
        expl_para = cell.add_paragraph()
        expl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        expl_run = expl_para.add_run(explanation)
        expl_run.font.size = Pt(9)
        expl_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def add_chart_placeholder(document, chart_spec, index: int) -> None:
    title = getattr(chart_spec, "title", None) or f"Grafico sugerido {index}"
    chart_type = str(getattr(chart_spec, "chart_type", "") or "")
    explanation = str(getattr(chart_spec, "plain_language_explanation", "") or "")
    description = str(getattr(chart_spec, "description", "") or "")

    caption = document.add_paragraph()
    caption.style = "Caption"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(f"Figura {index}. {title}")

    label = f"[Figura {index}: {title}]"
    _placeholder_box(document, label=label, chart_type=chart_type, explanation=explanation)

    if description:
        note = document.add_paragraph()
        note.style = "Small Text"
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.add_run(f"Finalidad: {description}")

    spacer = document.add_paragraph()
    spacer.style = "Normal"


def add_future_chart_note(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "Small Text"
    paragraph.add_run(
        "Nota. En esta fase el informe puede incluir imagenes exportadas del editor visual o placeholders "
        "visuales cuando no haya recursos PNG disponibles para Word."
    )


def add_chart_image_caption(document, title: str, index: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"Figura {index}. {title}")


def add_chart_image_note(document, note: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "Small Text"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(note)


def add_chart_image(document, chart_image, index: int) -> None:
    title = getattr(chart_image, "title", None) or f"Grafico exportado {index}"
    file_path_str = getattr(chart_image, "file_path", None)

    if not file_path_str or not Path(file_path_str).is_file():
        chart_type = str(getattr(chart_image, "chart_type", "") or "")
        add_chart_placeholder(
            document,
            type("_Stub", (), {"title": title, "chart_type": chart_type, "plain_language_explanation": "", "description": "Imagen no disponible en disco."})(),
            index,
        )
        return

    add_chart_image_caption(document, title, index)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(file_path_str, width=Inches(6.0))
    add_chart_image_note(
        document,
        "Nota. La imagen corresponde a una exportacion visual generada desde el editor de graficos de Colmena.",
    )
    spacer = document.add_paragraph()
    spacer.style = "Normal"


def add_chart_image_or_placeholder(document, chart_spec_or_image, index: int) -> None:
    file_path = getattr(chart_spec_or_image, "file_path", None)
    mime_type = getattr(chart_spec_or_image, "mime_type", "")
    if file_path and mime_type == "image/png":
        add_chart_image(document, chart_spec_or_image, index)
        return
    add_chart_placeholder(document, chart_spec_or_image, index)
