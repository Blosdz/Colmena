from __future__ import annotations

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def setup_margins(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def setup_styles(document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for style_name, size, bold in [("Heading 1", 14, True), ("Heading 2", 13, True), ("Heading 3", 12, True)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.space_before = Pt(6)

    for name, size in {"Table Text": 10, "Table Note": 10, "Caption": 11, "Small Text": 10}.items():
        style = document.styles[name] if name in document.styles else document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing = 1.0 if name != "Table Text" else 1.15
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.space_before = Pt(0)
        if name == "Caption":
            style.font.italic = True


def setup_document_defaults(document) -> None:
    for section in document.sections:
        setup_margins(section)
    setup_styles(document)
    document.core_properties.author = "Colmena"
    document.core_properties.company = "Colmena"
    document.core_properties.comments = "Documento academico generado automaticamente por Colmena."


def apply_heading_style(paragraph, level: int) -> None:
    paragraph.style = f"Heading {min(max(level, 1), 3)}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_body_style(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_table_style(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Text"
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)


def set_cell_text(cell, text: str, bold: bool = False, italic: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Table Text"
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
