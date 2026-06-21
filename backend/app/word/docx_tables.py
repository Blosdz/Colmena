from __future__ import annotations

from collections.abc import Iterable

from docx.shared import Cm, Pt

from app.schemas.apa_table import ApaTableRead
from app.word.docx_styles import apply_table_style, set_cell_text, set_repeat_table_header


def _safe_cell_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value)
    lowered = text.strip().lower()
    if lowered in {"none", "nan"}:
        return ""
    return text


def add_table_title(document, table_number: int, title: str) -> None:
    number_paragraph = document.add_paragraph()
    number_paragraph.style = "Caption"
    number_paragraph.add_run(f"Tabla {table_number}").bold = True

    title_paragraph = document.add_paragraph()
    title_paragraph.style = "Caption"
    title_paragraph.add_run(title).italic = True


def set_column_widths_safe(table) -> None:
    if not table.columns:
        return
    width = 16.0 / max(len(table.columns), 1)
    for column in table.columns:
        for cell in column.cells:
            cell.width = Cm(width)


def format_docx_table(table) -> None:
    apply_table_style(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Text"
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)


def add_table_body(document, columns, rows):
    table = document.add_table(rows=max(len(rows), 1) + 1, cols=max(len(columns), 1))
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, column in enumerate(columns):
        set_cell_text(header_row.cells[index], _safe_cell_text(column.label), bold=True)
    for row_index, row in enumerate(rows, start=1):
        for column_index, cell in enumerate(row.cells):
            set_cell_text(table.rows[row_index].cells[column_index], _safe_cell_text(cell.value))
    format_docx_table(table)
    set_column_widths_safe(table)
    return table


def add_table_notes(document, notes: Iterable) -> None:
    for index, note in enumerate(notes):
        paragraph = document.add_paragraph()
        paragraph.style = "Table Note"
        prefix = "Nota. " if index == 0 else ""
        paragraph.add_run(prefix + _safe_cell_text(note.text))


def add_apa_table_to_docx(document, apa_table: ApaTableRead, table_number: int) -> None:
    add_table_title(document, table_number, apa_table.title)
    add_table_body(document, apa_table.columns, apa_table.rows)
    add_table_notes(document, apa_table.notes)
