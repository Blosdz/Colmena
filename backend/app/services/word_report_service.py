from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from typing import Any

from docx import Document
from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.analysis_run import AnalysisRun
from app.models.export_artifact import ExportArtifact
from app.models.form import Form
from app.models.form_response import FormResponse
from app.schemas.apa_table import ApaTableRead, ApaTableRequest
from app.schemas.chart_image import ChartImageRead
from app.schemas.word_report import (
    AvailableWordAnalysisRunRead,
    WordReportGenerateRequest,
    WordReportOptionsRead,
    WordReportRead,
    WordReportSectionRead,
)
from app.services.advanced_scoring_service import AdvancedScoringService
from app.services.analysis_orchestrator_service import AnalysisOrchestratorService
from app.services.apa_table_service import ApaTableService
from app.services.chart_service import ChartService
from app.services.dataset_service import DatasetService
from app.word.docx_builder import add_page_break, create_document
from app.word.docx_export import (
    build_report_file_name,
    create_export_artifact_for_docx,
    ensure_exports_dir,
    get_file_size,
    save_docx,
)
from app.word.docx_sections import (
    build_appendix_section,
    build_categorical_association_section,
    build_chart_placeholder_section,
    build_conclusion_section,
    build_correlation_section,
    build_cover_section,
    build_dataset_section,
    build_descriptive_section,
    build_group_comparison_section,
    build_normality_section,
    build_project_section,
    build_scoring_section,
)


class WordReportService:
    def __init__(self, db: Session):
        self.db = db
        self.settings = get_settings()
        self.exports_dir = ensure_exports_dir(self.settings.backend_dir / "data" / "exports")
        self.dataset_service = DatasetService(db)
        self.apa_table_service = ApaTableService(db)
        self.chart_service = ChartService(db)
        self.advanced_scoring_service = AdvancedScoringService(db)
        self.analysis_orchestrator_service = AnalysisOrchestratorService(db)

    def _get_form(self, form_id: str) -> Form:
        return self.dataset_service._get_form(form_id)

    def _get_analysis_run(self, form_id: str, analysis_run_id: str) -> AnalysisRun:
        analysis_run = self.db.scalar(
            select(AnalysisRun).where(AnalysisRun.id == analysis_run_id, AnalysisRun.form_id == form_id)
        )
        if analysis_run is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="AnalysisRun not found for form")
        return analysis_run

    def _list_latest_runs(self, form_id: str, limit: int = 30) -> list[AnalysisRun]:
        return list(
            self.db.scalars(
                select(AnalysisRun)
                .where(AnalysisRun.form_id == form_id)
                .order_by(AnalysisRun.created_at.desc())
                .limit(limit)
            ).all()
        )

    def _list_word_artifacts(self, form_id: str) -> list[ExportArtifact]:
        return list(
            self.db.scalars(
                select(ExportArtifact)
                .where(
                    ExportArtifact.form_id == form_id,
                    ExportArtifact.artifact_type == "word_report_docx",
                )
                .order_by(ExportArtifact.created_at.desc())
            ).all()
        )

    def _list_chart_image_artifacts(self, form_id: str) -> list[ExportArtifact]:
        return list(
            self.db.scalars(
                select(ExportArtifact)
                .where(
                    ExportArtifact.form_id == form_id,
                    ExportArtifact.artifact_type.in_(("chart_image_png", "chart_image_svg")),
                )
                .order_by(ExportArtifact.created_at.desc())
            ).all()
        )

    def _safe_text(self, value: Any, fallback: str = "No especificado") -> str:
        if value is None:
            return fallback
        text = str(value).strip()
        if not text or text.lower() in {"none", "nan"}:
            return fallback
        return text

    def _artifact_to_read(self, artifact: ExportArtifact) -> WordReportRead:
        metadata = artifact.metadata_json if isinstance(artifact.metadata_json, dict) else {}
        sections = [
            WordReportSectionRead(**section)
            for section in metadata.get("sections", [])
            if isinstance(section, dict)
        ]
        return WordReportRead(
            artifact_id=artifact.id,
            form_id=artifact.form_id or "",
            project_id=artifact.project_id,
            report_type=str(metadata.get("report_type", "full_form_report")),
            file_name=artifact.file_name,
            file_path=(self.settings.backend_dir / artifact.file_path).as_posix(),
            mime_type=artifact.mime_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size_bytes=artifact.file_size_bytes or 0,
            sections=sections,
            table_count=int(metadata.get("table_count", 0)),
            chart_image_count=int(metadata.get("chart_image_count", 0)),
            chart_placeholder_count=int(metadata.get("chart_placeholder_count", 0)),
            chart_image_artifact_ids=[str(item) for item in metadata.get("chart_image_artifact_ids", []) if item],
            analysis_run_count=int(metadata.get("analysis_run_count", 0)),
            warnings=[str(item) for item in metadata.get("warnings", []) if item],
            created_at=artifact.created_at,
        )

    def _chart_image_to_read(self, artifact: ExportArtifact) -> ChartImageRead:
        metadata = artifact.metadata_json if isinstance(artifact.metadata_json, dict) else {}
        return ChartImageRead(
            artifact_id=artifact.id,
            form_id=artifact.form_id or "",
            project_id=artifact.project_id,
            chart_id=metadata.get("chart_id"),
            chart_type=metadata.get("chart_type"),
            title=metadata.get("title"),
            format=str(metadata.get("format", "png")),
            file_name=artifact.file_name,
            file_path=(self.settings.backend_dir / artifact.file_path).as_posix(),
            mime_type=artifact.mime_type or "image/png",
            file_size_bytes=artifact.file_size_bytes or 0,
            created_at=artifact.created_at,
            metadata_json=metadata,
        )

    def _append_section(
        self,
        sections: list[WordReportSectionRead],
        *,
        section_key: str,
        title: str,
        included: bool,
        summary: str,
        warnings: list[str] | None = None,
    ) -> None:
        sections.append(
            WordReportSectionRead(
                section_key=section_key,
                title=title,
                included=included,
                summary=summary,
                warnings=list(dict.fromkeys(warnings or [])),
            )
        )

    def _response_status_counts(self, form_id: str) -> dict[str, int]:
        counts: dict[str, int] = defaultdict(int)
        for response in self.db.scalars(select(FormResponse).where(FormResponse.form_id == form_id)).all():
            counts[response.status] += 1
        return counts

    def _project_overview_lines(self, form: Form, include_methodology_summary: bool) -> list[str]:
        project = form.project
        lines = [
            f"Proyecto: {self._safe_text(project.title)}.",
            f"Formulario: {self._safe_text(form.title)}.",
            f"Institucion: {self._safe_text(project.institution)}.",
            f"Facultad: {self._safe_text(project.faculty)}.",
            f"Carrera: {self._safe_text(project.career)}.",
            f"Asesor: {self._safe_text(project.advisor_name)}.",
        ]
        if include_methodology_summary:
            lines.extend(
                [
                    f"Tipo de investigacion: {self._safe_text(project.research_type)}.",
                    f"Diseno: {self._safe_text(project.design_type)}.",
                    f"Enfoque: {self._safe_text(project.approach)}.",
                    f"Poblacion: {self._safe_text(project.population_description)}.",
                    f"Muestra planificada: {self._safe_text(project.sample_size_planned)}.",
                ]
            )
        return lines

    def _dataset_summary_lines(self, form: Form, include_discarded: bool) -> tuple[list[str], list[str]]:
        summary = self.analysis_orchestrator_service.get_analysis_summary(form.id)
        option_summary = self.analysis_orchestrator_service.list_analysis_options(form.id)
        response_counts = self._response_status_counts(form.id)
        instrument_count = len([item for item in form.instruments if item.deleted_at is None])
        dimension_count = sum(
            1
            for instrument in form.instruments
            if instrument.deleted_at is None
            for dimension in instrument.dimensions
            if dimension.deleted_at is None
        )
        lines = [
            f"Total de respuestas: {summary.total_responses}.",
            f"Respuestas incluidas: {summary.included_responses}.",
            f"Respuestas descartadas: {response_counts.get('discarded', 0)}.",
            f"Porcentaje general de datos perdidos: {self._safe_text(summary.data_quality.get('missing_percent'), '0')}%.",
            f"Total de preguntas: {self._safe_text(summary.data_quality.get('total_questions'), '0')}.",
            f"Variables principales detectadas: {len(option_summary.available_targets['all'])}.",
            f"Instrumentos activos: {instrument_count}.",
            f"Dimensiones activas: {dimension_count}.",
            f"Respuestas descartadas incluidas en esta corrida: {'Si' if include_discarded else 'No'}.",
        ]
        return lines, list(summary.warnings)

    def _ensure_clean_visible_text(self, document: Document) -> None:
        visible = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if "None" in visible or "nan" in visible.lower():
            raise RuntimeError("Generated DOCX contains forbidden placeholder text")

    def _latest_runs_by_type(self, form_id: str) -> dict[str, AnalysisRun]:
        latest: dict[str, AnalysisRun] = {}
        for run in self._list_latest_runs(form_id, limit=50):
            latest.setdefault(run.analysis_type, run)
        return latest

    def _resolve_report_runs(self, form_id: str, request: WordReportGenerateRequest) -> tuple[list[AnalysisRun], AnalysisRun | None]:
        analysis_runs = [self._get_analysis_run(form_id, run_id) for run_id in request.analysis_run_ids or []]
        orchestrated_run = (
            self._get_analysis_run(form_id, request.orchestrated_analysis_run_id)
            if request.orchestrated_analysis_run_id
            else None
        )
        latest = self._latest_runs_by_type(form_id)
        if request.report_type == "orchestrated_report" and orchestrated_run is None:
            orchestrated_run = latest.get("orchestrated_analysis")
        if request.report_type == "inferential_report" and not analysis_runs:
            for key in ["advanced_scoring", "correlation", "group_comparison", "categorical_association", "normality"]:
                if key in latest:
                    analysis_runs.append(latest[key])
        if request.report_type == "full_form_report" and not analysis_runs:
            for key in ["advanced_scoring", "correlation", "group_comparison", "categorical_association"]:
                if key in latest:
                    analysis_runs.append(latest[key])
        return analysis_runs, orchestrated_run

    def build_report_tables(
        self,
        form_id: str,
        request: WordReportGenerateRequest,
        analysis_runs: list[AnalysisRun],
        orchestrated_run: AnalysisRun | None,
    ) -> dict[str, list[ApaTableRead]]:
        groups: dict[str, list[ApaTableRead]] = {
            "descriptive": [],
            "normality": [],
            "scoring": [],
            "correlation": [],
            "group_comparison": [],
            "categorical_association": [],
            "orchestrated": [],
        }
        if request.report_type in {"descriptive_report", "full_form_report", "inferential_report"}:
            groups["descriptive"].append(
                self.apa_table_service.generate_apa_table(
                    form_id,
                    ApaTableRequest(
                        table_type="descriptives",
                        source_type="live",
                        form_id=form_id,
                        decimals=request.decimals,
                        include_discarded=request.include_discarded,
                        score_aggregation=request.score_aggregation,
                    ),
                )
            )
        if request.report_type in {"descriptive_report", "full_form_report"}:
            groups["descriptive"].append(
                self.apa_table_service.generate_apa_table(
                    form_id,
                    ApaTableRequest(
                        table_type="frequencies",
                        source_type="live",
                        form_id=form_id,
                        decimals=request.decimals,
                        include_discarded=request.include_discarded,
                        score_aggregation=request.score_aggregation,
                    ),
                )
            )
        if request.report_type in {"inferential_report", "full_form_report"}:
            groups["normality"].append(
                self.apa_table_service.generate_apa_table(
                    form_id,
                    ApaTableRequest(
                        table_type="normality",
                        source_type="live",
                        form_id=form_id,
                        decimals=request.decimals,
                        include_discarded=request.include_discarded,
                        score_aggregation=request.score_aggregation,
                    ),
                )
            )
        if self.advanced_scoring_service._get_scoring_configs(form_id):
            for table_type in ["scoring_summary", "score_band_distribution", "control_scale_flags"]:
                table = self.apa_table_service.generate_apa_table(
                    form_id,
                    ApaTableRequest(
                        table_type=table_type,
                        source_type="live",
                        form_id=form_id,
                        decimals=request.decimals,
                        include_discarded=request.include_discarded,
                        score_aggregation=request.score_aggregation,
                    ),
                )
                if table.rows:
                    groups["scoring"].append(table)
        if orchestrated_run is not None:
            groups["orchestrated"].extend(
                self.apa_table_service.generate_from_orchestrated_analysis(
                    form_id,
                    orchestrated_run.id,
                    decimals=request.decimals,
                ).tables
            )
        for run in analysis_runs:
            batch = self.apa_table_service.generate_from_analysis_run(
                form_id,
                run.id,
                decimals=request.decimals,
                include_discarded=request.include_discarded,
                score_aggregation=request.score_aggregation,
            )
            for table in batch.tables:
                if table.table_type in {"scoring_summary", "score_band_distribution", "control_scale_flags"}:
                    groups["scoring"].append(table)
                elif table.table_type in {"correlation", "correlation_matrix"}:
                    groups["correlation"].append(table)
                elif table.table_type == "group_comparison":
                    groups["group_comparison"].append(table)
                elif table.table_type == "categorical_association":
                    groups["categorical_association"].append(table)
                elif table.table_type == "normality":
                    groups["normality"].append(table)
                else:
                    groups["descriptive"].append(table)
        return groups

    def build_chart_placeholders(
        self,
        form_id: str,
        request: WordReportGenerateRequest,
        analysis_runs: list[AnalysisRun],
        orchestrated_run: AnalysisRun | None,
    ) -> list[Any]:
        if not request.include_charts_placeholders:
            return []
        if orchestrated_run is not None:
            return self.chart_service.generate_from_orchestrated_run(
                form_id,
                orchestrated_run.id,
                theme="academic_light",
                decimals=request.decimals,
                include_discarded=request.include_discarded,
                score_aggregation=request.score_aggregation,
            ).charts
        if analysis_runs:
            charts: list[Any] = []
            for run in analysis_runs:
                charts.extend(
                    self.chart_service.generate_from_analysis_run(
                        form_id,
                        run.id,
                        theme="academic_light",
                        decimals=request.decimals,
                        include_discarded=request.include_discarded,
                        score_aggregation=request.score_aggregation,
                    ).charts
                )
            return charts[:8]
        return self.chart_service.generate_recommended_charts(
            form_id,
            theme="academic_light",
            decimals=request.decimals,
            include_discarded=request.include_discarded,
            score_aggregation=request.score_aggregation,
            max_charts=6,
        ).charts

    def _resolve_chart_images(self, form_id: str, request: WordReportGenerateRequest) -> tuple[list[ChartImageRead], list[str]]:
        warnings: list[str] = []
        if not request.include_chart_images:
            return [], warnings

        artifacts = self._list_chart_image_artifacts(form_id)
        by_id = {artifact.id: artifact for artifact in artifacts}

        if request.chart_image_mode == "placeholders_only":
            return [], warnings

        if request.chart_image_mode == "selected_images_only":
            selected_ids = request.chart_image_artifact_ids or []
            if not selected_ids:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="chart_image_artifact_ids is required for selected_images_only",
                )
            missing = [artifact_id for artifact_id in selected_ids if artifact_id not in by_id]
            if missing:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="One or more chart images do not belong to the form",
                )
            selected = [self._chart_image_to_read(by_id[artifact_id]) for artifact_id in selected_ids]
            png_selected = [item for item in selected if item.format == "png"]
            if len(png_selected) < len(selected):
                warnings.append("svg_not_supported_in_word")
            return png_selected, warnings

        images = [self._chart_image_to_read(artifact) for artifact in artifacts if artifact.metadata_json]
        png_images = [item for item in images if item.format == "png"]
        if images and not png_images:
            warnings.append("svg_not_supported_in_word")
        return png_images, warnings

    def build_report_warnings(
        self,
        form: Form,
        request: WordReportGenerateRequest,
        analysis_runs: list[AnalysisRun],
        tables: dict[str, list[ApaTableRead]],
        chart_image_warnings: list[str],
    ) -> list[str]:
        warnings = list(self.analysis_orchestrator_service.get_analysis_summary(form.id).warnings)
        if not request.include_discarded:
            warnings.append("discarded_responses_excluded")
        if not analysis_runs and request.report_type in {"inferential_report", "full_form_report"}:
            warnings.append("no_inferential_runs")
        for table_list in tables.values():
            for table in table_list:
                warnings.extend(table.warnings)
        warnings.extend(chart_image_warnings)
        return list(dict.fromkeys(warnings))

    def _run_interpretations(self, analysis_runs: list[AnalysisRun], analysis_type: str) -> list[str]:
        texts: list[str] = []
        for run in analysis_runs:
            if run.analysis_type != analysis_type:
                continue
            result = run.result_json if isinstance(run.result_json, dict) else {}
            text = result.get("interpretation") or result.get("academic_interpretation")
            if text:
                texts.append(self._safe_text(text, fallback=""))
        return [item for item in dict.fromkeys(texts) if item]

    def build_conclusions(
        self,
        form_id: str,
        analysis_runs: list[AnalysisRun],
        orchestrated_run: AnalysisRun | None,
        warnings: list[str],
    ) -> list[str]:
        lines: list[str] = []
        if orchestrated_run is not None:
            result = orchestrated_run.result_json if isinstance(orchestrated_run.result_json, dict) else {}
            for key in ["executive_summary", "academic_interpretation", "plain_language_explanation"]:
                value = result.get(key)
                if value:
                    lines.append(self._safe_text(value, fallback=""))
        for run in analysis_runs:
            result = run.result_json if isinstance(run.result_json, dict) else {}
            text = result.get("interpretation") or result.get("academic_interpretation")
            if text:
                lines.append(self._safe_text(text, fallback=""))
        if not lines:
            summary = self.analysis_orchestrator_service.get_analysis_summary(form_id)
            if summary.total_responses == 0:
                lines.append("No se dispone de respuestas suficientes para una conclusion estadistica robusta.")
            else:
                lines.append("La base de datos permite un resumen descriptivo inicial, pero aun no cuenta con suficientes resultados inferenciales reutilizables.")
        if "high_missingness" in warnings:
            lines.append("Se recomienda revisar la perdida de datos antes de sostener conclusiones firmes.")
        if "posthoc_required" in warnings:
            lines.append("Si se detectaron diferencias globales entre mas de dos grupos, sera necesario ejecutar comparaciones post hoc en una fase posterior.")
        return [item for item in dict.fromkeys(lines) if item]

    def build_appendix_lines(
        self,
        analysis_runs: list[AnalysisRun],
        orchestrated_run: AnalysisRun | None,
        request: WordReportGenerateRequest,
    ) -> list[str]:
        lines = [
            f"Decimales configurados: {request.decimals}.",
            f"Agregacion de puntajes: {request.score_aggregation}.",
            f"Incluir respuestas descartadas: {'Si' if request.include_discarded else 'No'}.",
        ]
        for run in analysis_runs:
            lines.append(f"AnalysisRun utilizado: {run.id} ({run.analysis_type}).")
        if orchestrated_run is not None:
            lines.append(f"AnalysisRun orquestado utilizado: {orchestrated_run.id}.")
        return lines

    def build_report_sections(
        self,
        document: Document,
        form: Form,
        request: WordReportGenerateRequest,
        analysis_runs: list[AnalysisRun],
        orchestrated_run: AnalysisRun | None,
        tables: dict[str, list[ApaTableRead]],
        charts: list[Any],
        chart_images: list[ChartImageRead],
        warnings: list[str],
    ) -> list[WordReportSectionRead]:
        sections: list[WordReportSectionRead] = []
        generated_at = datetime.now()
        title = request.title or "Informe de resultados estadisticos"
        subtitle = request.subtitle or "Reporte generado por Colmena"
        table_number = 1

        if request.include_cover:
            build_cover_section(document, title=title, subtitle=subtitle, project=form.project, generated_at=generated_at)
            self._append_section(
                sections,
                section_key="cover",
                title="Portada",
                included=True,
                summary="Se incluyo una portada simple con datos del proyecto y fecha de generacion.",
            )
            add_page_break(document)

        build_project_section(document, self._project_overview_lines(form, request.include_methodology_summary))
        self._append_section(
            sections,
            section_key="project_overview",
            title="Datos generales del estudio",
            included=True,
            summary="Se resumieron los datos del proyecto, el formulario y la informacion metodologica disponible.",
        )

        dataset_lines, dataset_warnings = self._dataset_summary_lines(form, request.include_discarded)
        build_dataset_section(document, dataset_lines)
        self._append_section(
            sections,
            section_key="dataset_summary",
            title="Resumen de la base de datos",
            included=True,
            summary="Se documentaron respuestas, completitud y estructura general de la base.",
            warnings=dataset_warnings,
        )

        if tables["descriptive"]:
            table_number = build_descriptive_section(
                document,
                "Se presentan frecuencias y descriptivos iniciales para describir la muestra y las variables principales del formulario.",
                tables["descriptive"],
                table_number,
            )
            self._append_section(
                sections,
                section_key="descriptive_results",
                title="Resultados descriptivos",
                included=True,
                summary="Se integraron tablas APA de frecuencias y descriptivos numericos.",
                warnings=[warning for table in tables["descriptive"] for warning in table.warnings],
            )

        if tables["normality"]:
            table_number = build_normality_section(
                document,
                "Las pruebas de normalidad se reportan segun el nivel de significancia configurado y deben interpretarse con prudencia metodologica.",
                tables["normality"],
                table_number,
            )
            self._append_section(
                sections,
                section_key="normality_results",
                title="Pruebas de normalidad",
                included=True,
                summary="Se incluyeron resultados de normalidad para los targets numericos aplicables.",
                warnings=[warning for table in tables["normality"] for warning in table.warnings],
            )

        if tables["scoring"]:
            table_number = build_scoring_section(
                document,
                "Se presentan puntajes, distribuciones por baremo y resultados de escalas de control cuando existen configuraciones avanzadas para el formulario.",
                tables["scoring"],
                table_number,
            )
            self._append_section(
                sections,
                section_key="scoring_results",
                title="Resultados por baremo",
                included=True,
                summary="Se integraron tablas APA de scoring avanzado, distribucion por niveles y flags de control.",
                warnings=[warning for table in tables["scoring"] for warning in table.warnings],
            )

        if tables["correlation"]:
            table_number = build_correlation_section(
                document,
                "Se reportan asociaciones entre variables segun el metodo estadistico aplicado en cada corrida reutilizada.",
                tables["correlation"],
                self._run_interpretations(analysis_runs, "correlation"),
                table_number,
            )
            self._append_section(
                sections,
                section_key="correlation_results",
                title="Correlaciones",
                included=True,
                summary="Se documentaron correlaciones y sus interpretaciones academicas.",
                warnings=[warning for table in tables["correlation"] for warning in table.warnings],
            )

        if tables["group_comparison"]:
            table_number = build_group_comparison_section(
                document,
                "Se presentan comparaciones entre grupos con sus estadisticos, niveles de significancia y tamanos del efecto cuando estuvieron disponibles.",
                tables["group_comparison"],
                self._run_interpretations(analysis_runs, "group_comparison"),
                table_number,
            )
            self._append_section(
                sections,
                section_key="group_comparison_results",
                title="Comparaciones entre grupos",
                included=True,
                summary="Se incluyeron resultados inferenciales de comparacion entre grupos.",
                warnings=[warning for table in tables["group_comparison"] for warning in table.warnings],
            )

        if tables["categorical_association"]:
            table_number = build_categorical_association_section(
                document,
                "Se reportan pruebas de asociacion categorica con sus advertencias metodologicas correspondientes.",
                tables["categorical_association"],
                self._run_interpretations(analysis_runs, "categorical_association"),
                table_number,
            )
            self._append_section(
                sections,
                section_key="categorical_association_results",
                title="Asociacion categorica",
                included=True,
                summary="Se incorporaron tablas y pruebas de asociacion categorica.",
                warnings=[warning for table in tables["categorical_association"] for warning in table.warnings],
            )

        if orchestrated_run is not None and tables["orchestrated"]:
            table_number = build_descriptive_section(
                document,
                "Se materializaron ademas bloques consolidados provenientes de una corrida orquestada previa.",
                tables["orchestrated"],
                table_number,
            )
            self._append_section(
                sections,
                section_key="orchestrated_results",
                title="Resultados orquestados consolidados",
                included=True,
                summary="Se transformaron los bloques APA del orquestador en tablas reutilizables dentro del documento.",
                warnings=[warning for table in tables["orchestrated"] for warning in table.warnings],
            )

        if request.include_charts_placeholders or chart_images:
            chart_visuals = chart_images if chart_images else charts
            image_count, placeholder_count = build_chart_placeholder_section(document, chart_visuals)
            self._append_section(
                sections,
                section_key="chart_placeholders",
                title="Graficos sugeridos",
                included=True,
                summary=(
                    "Se integraron graficos exportados desde el editor visual."
                    if image_count
                    else "Se anadieron placeholders para los graficos sugeridos en fases posteriores."
                ),
                warnings=[
                    *[warning for chart in charts for warning in getattr(chart, "warnings", [])],
                    *warnings,
                ],
            )

        build_conclusion_section(document, self.build_conclusions(form.id, analysis_runs, orchestrated_run, warnings))
        self._append_section(
            sections,
            section_key="conclusions",
            title="Conclusiones estadisticas preliminares",
            included=True,
            summary="Se sintetizaron hallazgos principales, cautelas metodologicas y proximos pasos analiticos.",
            warnings=warnings,
        )

        if request.include_technical_appendix:
            build_appendix_section(document, self.build_appendix_lines(analysis_runs, orchestrated_run, request))
            self._append_section(
                sections,
                section_key="technical_appendix",
                title="Anexo tecnico",
                included=True,
                summary="Se anadio un anexo con identificadores de corridas y parametros tecnicos.",
            )

        return sections

    def _save_artifact(
        self,
        form: Form,
        document: Document,
        report_type: str,
        sections: list[WordReportSectionRead],
        table_count: int,
        chart_image_count: int,
        chart_placeholder_count: int,
        chart_image_artifact_ids: list[str],
        analysis_run_count: int,
        warnings: list[str],
    ) -> WordReportRead:
        file_name = build_report_file_name(form.id, report_type)
        output_path = self.exports_dir / file_name
        save_docx(document, output_path)
        file_size_bytes = get_file_size(output_path)
        relative_path = output_path.relative_to(self.settings.backend_dir).as_posix()

        artifact = create_export_artifact_for_docx(
            form=form,
            file_name=file_name,
            relative_path=relative_path,
            file_size_bytes=file_size_bytes,
            metadata_json={
                "report_type": report_type,
                "sections": [section.model_dump() for section in sections],
                "table_count": table_count,
                "chart_image_count": chart_image_count,
                "chart_placeholder_count": chart_placeholder_count,
                "chart_image_artifact_ids": chart_image_artifact_ids,
                "analysis_run_count": analysis_run_count,
                "warnings": warnings,
            },
        )
        self.db.add(artifact)
        self.db.commit()
        self.db.refresh(artifact)
        return self._artifact_to_read(artifact)

    def generate_word_report(self, form_id: str, request: WordReportGenerateRequest) -> WordReportRead:
        form = self._get_form(form_id)
        if request.source_type == "analysis_run" and not request.analysis_run_ids:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="analysis_run_ids is required for source_type analysis_run")
        if request.source_type == "orchestrated" and not request.orchestrated_analysis_run_id:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="orchestrated_analysis_run_id is required for source_type orchestrated")

        analysis_runs, orchestrated_run = self._resolve_report_runs(form_id, request)
        if request.report_type == "orchestrated_report" and orchestrated_run is None:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="No orchestrated analysis run available")

        tables = self.build_report_tables(form_id, request, analysis_runs, orchestrated_run)
        charts = self.build_chart_placeholders(form_id, request, analysis_runs, orchestrated_run)
        chart_images, chart_image_warnings = self._resolve_chart_images(form_id, request)
        warnings = self.build_report_warnings(form, request, analysis_runs, tables, chart_image_warnings)

        document = create_document()
        sections = self.build_report_sections(
            document,
            form,
            request,
            analysis_runs,
            orchestrated_run,
            tables,
            charts,
            chart_images,
            warnings,
        )
        self._ensure_clean_visible_text(document)

        table_count = sum(len(group) for group in tables.values())
        analysis_run_count = len(analysis_runs) + (1 if orchestrated_run is not None else 0)
        chart_image_count = len(chart_images)
        chart_placeholder_count = len(charts) if (not chart_images and request.include_charts_placeholders) else 0
        return self._save_artifact(
            form,
            document,
            request.report_type,
            sections,
            table_count,
            chart_image_count,
            chart_placeholder_count,
            [item.artifact_id for item in chart_images],
            analysis_run_count,
            warnings,
        )

    def generate_from_analysis_run_report(self, form_id: str, analysis_run_id: str) -> WordReportRead:
        analysis_run = self._get_analysis_run(form_id, analysis_run_id)
        report_type = "descriptive_report" if analysis_run.analysis_type == "descriptive" else "inferential_report"
        return self.generate_word_report(
            form_id,
            WordReportGenerateRequest(
                report_type=report_type,
                source_type="analysis_run",
                analysis_run_ids=[analysis_run_id],
                include_charts_placeholders=True,
                include_chart_images=True,
                include_plain_language_explanations=True,
                include_cover=True,
                include_methodology_summary=True,
            ),
        )

    def generate_from_orchestrated_run_report(self, form_id: str, analysis_run_id: str) -> WordReportRead:
        return self.generate_word_report(
            form_id,
            WordReportGenerateRequest(
                report_type="orchestrated_report",
                source_type="orchestrated",
                orchestrated_analysis_run_id=analysis_run_id,
                include_charts_placeholders=True,
                include_chart_images=True,
                include_plain_language_explanations=True,
                include_cover=True,
                include_methodology_summary=True,
            ),
        )

    def get_word_report_options(self, form_id: str) -> WordReportOptionsRead:
        form = self._get_form(form_id)
        runs = self._list_latest_runs(form_id)
        available_runs = [
            AvailableWordAnalysisRunRead(id=run.id, analysis_type=run.analysis_type, created_at=run.created_at)
            for run in runs
            if run.analysis_type in {"descriptive", "normality", "correlation", "group_comparison", "categorical_association"}
        ]
        orchestrated_runs = [
            AvailableWordAnalysisRunRead(id=run.id, analysis_type=run.analysis_type, created_at=run.created_at)
            for run in runs
            if run.analysis_type == "orchestrated_analysis"
        ]
        return WordReportOptionsRead(
            form_id=form.id,
            available_report_types=[
                "descriptive_report",
                "inferential_report",
                "orchestrated_report",
                "full_form_report",
            ],
            available_analysis_runs=available_runs,
            available_orchestrated_runs=orchestrated_runs,
            recommended_report="full_form_report" if form.responses else "descriptive_report",
            available_sections=[
                "cover",
                "project_overview",
                "dataset_summary",
                "descriptive_results",
                "normality_results",
                "scoring_results",
                "correlation_results",
                "group_comparison_results",
                "categorical_association_results",
                "chart_placeholders",
                "conclusions",
                "technical_appendix",
            ],
        )

    def list_word_reports(self, form_id: str) -> list[WordReportRead]:
        self._get_form(form_id)
        return [self._artifact_to_read(item) for item in self._list_word_artifacts(form_id)]
